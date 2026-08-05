from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.schemas.ai import ProjectWeeklyReportResult


SECTION_TITLES = [
    ("一、本周整体进展", "overall_summary"),
    ("二、项目进度", "progress_summary"),
]

LIST_SECTIONS = [
    ("三、已完成工作", "completed_work"),
    ("四、进行中任务", "ongoing_tasks"),
    ("五、风险与阻塞", "blocked_or_risky_items"),
    ("六、近期汇报摘要", "recent_reports_summary"),
    ("七、下周计划", "next_week_plan"),
    ("八、AI 管理建议", "management_suggestions"),
]


def build_weekly_report_docx(report: ProjectWeeklyReportResult) -> bytes:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    title = document.add_heading(report.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(f"项目名称：{report.project_name}")
    document.add_paragraph(f"项目编码：{report.project_code or '-'}")
    document.add_paragraph(f"统计周期：{report.period}")

    for title_text, field in SECTION_TITLES:
        document.add_heading(title_text, level=1)
        document.add_paragraph(getattr(report, field) or "暂无内容。")

    for title_text, field in LIST_SECTIONS:
        document.add_heading(title_text, level=1)
        values = getattr(report, field) or []
        if not values:
            document.add_paragraph("暂无内容。")
            continue
        for value in values:
            document.add_paragraph(str(value), style="List Bullet")

    document.add_paragraph()
    note = document.add_paragraph("说明：本周报由 DriveMind AI 基于项目、任务和工作汇报数据生成，供项目管理和复盘参考。")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
